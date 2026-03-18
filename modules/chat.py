import base64
import copy
import functools
import html
import json
import pprint
import re
import shutil
import threading
import time
from datetime import datetime
from functools import partial
from pathlib import Path

import markupsafe
import yaml
from jinja2.ext import loopcontrols
from jinja2.sandbox import ImmutableSandboxedEnvironment
from PIL import Image

import modules.shared as shared
from modules import utils
from modules.extensions import apply_extensions
from modules.html_generator import (
    chat_html_wrapper,
    convert_to_markdown,
    extract_thinking_block,
    make_thumbnail
)
from modules.image_utils import open_image_safely
from modules.logging_colors import logger
from modules.reasoning import THINKING_FORMATS
from modules.text_generation import (
    generate_reply,
    get_encoded_length,
    get_max_prompt_length
)
from modules.utils import (
    delete_file,
    get_available_characters,
    get_available_users,
    sanitize_filename,
    save_file
)
from modules.web_search import add_web_search_attachments

_history_file_lock = threading.Lock()


def strftime_now(format):
    return datetime.now().strftime(format)


def get_current_timestamp():
    """Returns the current time in 24-hour format"""
    return datetime.now().strftime('%b %d, %Y %H:%M')


def update_message_metadata(metadata_dict, role, index, **fields):
    """
    Updates or adds metadata fields for a specific message.

    Args:
        metadata_dict: The metadata dictionary
        role: The role (user, assistant, etc)
        index: The message index
        **fields: Arbitrary metadata fields to update/add
    """
    key = f"{role}_{index}"
    if key not in metadata_dict:
        metadata_dict[key] = {}

    # Update with provided fields
    for field_name, field_value in fields.items():
        metadata_dict[key][field_name] = field_value


jinja_env = ImmutableSandboxedEnvironment(
    trim_blocks=True,
    lstrip_blocks=True,
    extensions=[loopcontrols]
)


def custom_tojson(value, indent=None, ensure_ascii=True):
    return markupsafe.Markup(json.dumps(value, indent=indent, ensure_ascii=ensure_ascii))


jinja_env.filters["tojson"] = custom_tojson
jinja_env.globals["strftime_now"] = strftime_now


def _raise_exception(message):
    raise ValueError(message)


jinja_env.globals["raise_exception"] = _raise_exception

_template_cache = {}


def get_compiled_template(template_str):
    """Cache compiled Jinja2 templates keyed by their source string."""
    compiled = _template_cache.get(template_str)
    if compiled is None:
        compiled = jinja_env.from_string(template_str)
        _template_cache[template_str] = compiled

    return compiled


def str_presenter(dumper, data):
    """
    Copied from https://github.com/yaml/pyyaml/issues/240
    Makes pyyaml output prettier multiline strings.
    """

    if data.count('\n') > 0:
        return dumper.represent_scalar('tag:yaml.org,2002:str', data, style='|')

    return dumper.represent_scalar('tag:yaml.org,2002:str', data)


yaml.add_representer(str, str_presenter)
yaml.representer.SafeRepresenter.add_representer(str, str_presenter)


class _JsonDict(dict):
    """A dict that serializes as JSON when used in string concatenation.

    Some Jinja2 templates (Qwen, GLM) iterate arguments with .items(),
    requiring a dict.  Others (DeepSeek) concatenate arguments as a
    string, requiring JSON.  This class satisfies both.
    """

    def __str__(self):
        return json.dumps(self, ensure_ascii=False)

    def __add__(self, other):
        return str(self) + other

    def __radd__(self, other):
        return other + str(self)


def _deserialize_tool_call_arguments(tool_calls):
    """Convert tool_call arguments from JSON strings to _JsonDict.

    The OpenAI API spec sends arguments as a JSON string, but Jinja2
    templates may need a dict (.items()) or a string (concatenation).
    _JsonDict handles both transparently.
    """
    result = []
    for tc in tool_calls:
        tc = copy.copy(tc)
        func = tc.get('function', {})
        if isinstance(func, dict):
            func = dict(func)
            args = func.get('arguments')
            if isinstance(args, str):
                try:
                    func['arguments'] = _JsonDict(json.loads(args))
                except (json.JSONDecodeError, ValueError):
                    pass
            elif isinstance(args, dict) and not isinstance(args, _JsonDict):
                func['arguments'] = _JsonDict(args)
            tc['function'] = func
        result.append(tc)
    return result


def _expand_tool_sequence(tool_seq):
    """Expand a tool_sequence list into API messages.

    Returns a list of dicts (role: assistant with tool_calls, or role: tool).
    If any tool_call IDs are missing a matching tool result, a synthetic
    empty result is inserted so the prompt is never malformed.
    """
    messages = []
    expected_ids = []
    seen_ids = set()

    for item in tool_seq:
        if 'tool_calls' in item:
            deserialized = _deserialize_tool_call_arguments(item['tool_calls'])
            messages.append({
                "role": "assistant",
                "content": item.get('content', ''),
                "tool_calls": deserialized
            })
            for tc in item['tool_calls']:
                tc_id = tc.get('id', '')
                if tc_id:
                    expected_ids.append(tc_id)
        elif item.get('role') == 'tool':
            messages.append({
                "role": "tool",
                "content": item['content'],
                "tool_call_id": item.get('tool_call_id', '')
            })
            seen_ids.add(item.get('tool_call_id', ''))

    # Fill in synthetic results for any orphaned tool call IDs
    for tc_id in expected_ids:
        if tc_id not in seen_ids:
            messages.append({
                "role": "tool",
                "content": "",
                "tool_call_id": tc_id
            })

    return messages


def generate_chat_prompt(user_input, state, **kwargs):
    impersonate = kwargs.get('impersonate', False)
    _continue = kwargs.get('_continue', False)
    also_return_rows = kwargs.get('also_return_rows', False)
    history_data = kwargs.get('history', state['history'])
    history = history_data['internal']
    metadata = history_data.get('metadata', {})

    # Templates
    chat_template_str = state['chat_template_str']
    if state['mode'] != 'instruct':
        chat_template_str = replace_character_names(chat_template_str, state['name1'], state['name2'])

    instruction_template = get_compiled_template(state['instruction_template_str'])
    chat_template = get_compiled_template(chat_template_str)

    instruct_renderer = partial(
        instruction_template.render,
        builtin_tools=None,
        tools=state['tools'] if 'tools' in state else None,
        tools_in_user_message=False,
        add_generation_prompt=False,
        enable_thinking=state['enable_thinking'],
        reasoning_effort=state['reasoning_effort'],
        thinking_budget=-1 if state.get('enable_thinking', True) else 0,
        bos_token=shared.bos_token,
        eos_token=shared.eos_token,
    )

    chat_renderer = partial(
        chat_template.render,
        add_generation_prompt=False,
        name1=state['name1'],
        name2=state['name2'],
        user_bio=replace_character_names(state['user_bio'], state['name1'], state['name2']),
        tools=state['tools'] if 'tools' in state else None,
    )

    messages = []

    if state['mode'] == 'instruct':
        renderer = instruct_renderer
        if state['custom_system_message'].strip() != '':
            messages.append({"role": "system", "content": state['custom_system_message']})
    else:
        renderer = chat_renderer
        if state['context'].strip() != '' or state['user_bio'].strip() != '':
            context = replace_character_names(state['context'], state['name1'], state['name2'])
            messages.append({"role": "system", "content": context})

    insert_pos = len(messages)
    for i, entry in enumerate(reversed(history)):
        user_msg = entry[0].strip()
        assistant_msg = entry[1].strip()
        tool_msg = entry[2].strip() if len(entry) > 2 else ''
        entry_meta = entry[3] if len(entry) > 3 else {}

        row_idx = len(history) - i - 1

        if tool_msg:
            tool_message = {"role": "tool", "content": tool_msg}
            if "tool_call_id" in entry_meta:
                tool_message["tool_call_id"] = entry_meta["tool_call_id"]
            messages.insert(insert_pos, tool_message)

        if not assistant_msg and entry_meta.get('tool_calls'):
            # Assistant message with only tool_calls and no text content
            messages.insert(insert_pos, {"role": "assistant", "content": "", "tool_calls": _deserialize_tool_call_arguments(entry_meta['tool_calls'])})
        elif assistant_msg:
            # Handle GPT-OSS as a special case
            if '<|channel|>analysis<|message|>' in assistant_msg or '<|channel|>final<|message|>' in assistant_msg:
                thinking_content = ""
                final_content = ""

                # Extract analysis content if present
                if '<|channel|>analysis<|message|>' in assistant_msg:
                    parts = assistant_msg.split('<|channel|>analysis<|message|>', 1)
                    if len(parts) > 1:
                        # The content is everything after the tag
                        potential_content = parts[1]

                        # Now, find the end of this content block
                        analysis_end_tag = '<|end|>'
                        if analysis_end_tag in potential_content:
                            thinking_content = potential_content.split(analysis_end_tag, 1)[0].strip()
                        else:
                            # Fallback: if no <|end|> tag, stop at the start of the final channel if it exists
                            final_channel_tag = '<|channel|>final<|message|>'
                            if final_channel_tag in potential_content:
                                thinking_content = potential_content.split(final_channel_tag, 1)[0].strip()
                            else:
                                thinking_content = potential_content.strip()

                # Extract final content if present
                final_tag_to_find = '<|channel|>final<|message|>'
                if final_tag_to_find in assistant_msg:
                    parts = assistant_msg.split(final_tag_to_find, 1)
                    if len(parts) > 1:
                        # The content is everything after the tag
                        potential_content = parts[1]

                        # Now, find the end of this content block
                        final_end_tag = '<|end|>'
                        if final_end_tag in potential_content:
                            final_content = potential_content.split(final_end_tag, 1)[0].strip()
                        else:
                            final_content = potential_content.strip()

                # Insert as structured message
                msg_dict = {"role": "assistant", "content": final_content}
                if '<|channel|>analysis<|message|>' in assistant_msg:
                    msg_dict["thinking"] = thinking_content

                messages.insert(insert_pos, msg_dict)

            # Handle Seed-OSS
            elif '<seed:think>' in assistant_msg:
                thinking_content = ""
                final_content = assistant_msg

                # Extract thinking content if present
                if '<seed:think>' in assistant_msg:
                    parts = assistant_msg.split('<seed:think>', 1)
                    if len(parts) > 1:
                        potential_content = parts[1]
                        if '</seed:think>' in potential_content:
                            thinking_content = potential_content.split('</seed:think>', 1)[0].strip()
                            final_content = parts[0] + potential_content.split('</seed:think>', 1)[1]
                        else:
                            thinking_content = potential_content.strip()
                            final_content = parts[0]

                # Insert as structured message
                msg_dict = {"role": "assistant", "content": final_content.strip()}
                if thinking_content:
                    msg_dict["reasoning_content"] = thinking_content

                messages.insert(insert_pos, msg_dict)

            else:
                # Default case (used by all other models)
                messages.insert(insert_pos, {"role": "assistant", "content": assistant_msg})

            # Attach tool_calls metadata to the assistant message if present
            if entry_meta.get('tool_calls') and messages[insert_pos].get('role') == 'assistant':
                messages[insert_pos]['tool_calls'] = _deserialize_tool_call_arguments(entry_meta['tool_calls'])

        # Expand tool_sequence from metadata (inserted AFTER assistant so that
        # the final order is: user → tool_calls → tool_results → final_answer)
        meta_key = f"assistant_{row_idx}"
        tool_seq = metadata.get(meta_key, {}).get('tool_sequence', [])
        if tool_seq:
            for msg in reversed(_expand_tool_sequence(tool_seq)):
                messages.insert(insert_pos, msg)

        if entry_meta.get('role') == 'system':
            if user_msg:
                messages.insert(insert_pos, {"role": "system", "content": user_msg})
        elif user_msg not in ['', '<|BEGIN-VISIBLE-CHAT|>']:
            # Check for user message attachments in metadata
            user_key = f"user_{row_idx}"
            enhanced_user_msg = user_msg

            # Add attachment content if present AND if past attachments are enabled
            if user_key in metadata and "attachments" in metadata[user_key]:
                attachments_text = ""
                image_refs = ""

                for attachment in metadata[user_key]["attachments"]:
                    if attachment.get("type") == "image":
                        # Add image reference for multimodal models
                        image_refs += "<__media__>"
                    elif state.get('include_past_attachments', True):
                        # Handle text/PDF attachments
                        filename = attachment.get("name", "file")
                        content = attachment.get("content", "")
                        if attachment.get("type") == "text/html" and attachment.get("url"):
                            attachments_text += f"\nName: {filename}\nURL: {attachment['url']}\nContents:\n\n=====\n{content}\n=====\n\n"
                        else:
                            attachments_text += f"\nName: {filename}\nContents:\n\n=====\n{content}\n=====\n\n"

                if image_refs:
                    enhanced_user_msg = f"{image_refs}\n\n{enhanced_user_msg}"
                if attachments_text:
                    enhanced_user_msg += f"\n\nATTACHMENTS:\n{attachments_text}"

            messages.insert(insert_pos, {"role": "user", "content": enhanced_user_msg})

    # Handle the current user input
    user_input = user_input.strip()

    # Check if we have attachments
    if not (impersonate or _continue):
        has_attachments = False
        if len(history_data.get('metadata', {})) > 0:
            current_row_idx = len(history)
            user_key = f"user_{current_row_idx}"
            has_attachments = user_key in metadata and "attachments" in metadata[user_key]

        if user_input or has_attachments:
            # For the current user input being processed, check if we need to add attachments
            if len(history_data.get('metadata', {})) > 0:
                current_row_idx = len(history)
                user_key = f"user_{current_row_idx}"

                if user_key in metadata and "attachments" in metadata[user_key]:
                    attachments_text = ""
                    image_refs = ""

                    for attachment in metadata[user_key]["attachments"]:
                        if attachment.get("type") == "image":
                            image_refs += "<__media__>"
                        else:
                            filename = attachment.get("name", "file")
                            content = attachment.get("content", "")
                            if attachment.get("type") == "text/html" and attachment.get("url"):
                                attachments_text += f"\nName: {filename}\nURL: {attachment['url']}\nContents:\n\n=====\n{content}\n=====\n\n"
                            else:
                                attachments_text += f"\nName: {filename}\nContents:\n\n=====\n{content}\n=====\n\n"

                    if image_refs:
                        user_input = f"{image_refs}\n\n{user_input}"
                    if attachments_text:
                        user_input += f"\n\nATTACHMENTS:\n{attachments_text}"

            messages.append({"role": "user", "content": user_input})

        # Expand tool_sequence for the current entry (excluded from the
        # history loop during regenerate — needed so the model sees prior
        # tool calls and results when re-generating the final answer).
        current_tool_seq = metadata.get(f"assistant_{len(history)}", {}).get('tool_sequence', [])
        messages.extend(_expand_tool_sequence(current_tool_seq))

    if impersonate and state['mode'] != 'chat-instruct':
        messages.append({"role": "user", "content": "fake user message replace me"})

    def make_prompt(messages):
        last_message = messages[-1].copy()
        if _continue:
            if state['mode'] == 'chat-instruct':
                messages = messages[:-1]
            else:
                messages[-1]["content"] = "fake assistant message replace me"
                messages.append({"role": "assistant", "content": "this will get deleted"})

        if state['mode'] != 'chat-instruct':
            add_generation_prompt = (not _continue and not impersonate)
        else:
            add_generation_prompt = False

        prompt = renderer(
            messages=messages,
            add_generation_prompt=add_generation_prompt
        )

        if state['mode'] == 'chat-instruct':
            command = state['chat-instruct_command']
            command = command.replace('<|character|>', state['name2'] if not impersonate else state['name1'])
            command = command.replace('<|prompt|>', prompt)
            command = replace_character_names(command, state['name1'], state['name2'])

            outer_messages = []
            if state['custom_system_message'].strip() != '':
                outer_messages.append({"role": "system", "content": state['custom_system_message']})

            outer_messages.append({"role": "user", "content": command})
            if _continue:
                outer_messages.append(last_message.copy())
                outer_messages[-1]["content"] = "fake assistant message replace me"
                outer_messages.append({"role": "assistant", "content": "this will get deleted"})

            prompt = instruct_renderer(
                messages=outer_messages,
                add_generation_prompt=not _continue
            )

        if _continue:
            prompt = prompt.split("fake assistant message replace me", 1)[0]

            content = last_message.get("content", "")
            partial_thought = last_message.get("thinking", "") or last_message.get("reasoning_content", "")

            # Handle partial thinking blocks (GPT-OSS and Seed-OSS)
            if not content and partial_thought and partial_thought.strip():
                search_string = partial_thought.strip()
                index = prompt.rfind(search_string)
                if index != -1:
                    prompt = prompt[:index] + partial_thought
                else:
                    # Fallback if search fails: just append the thought
                    prompt += partial_thought
            else:
                # All other cases
                prompt += content

        if impersonate:
            prompt = prompt.split("fake user message replace me", 1)[0]
            prompt += user_input

        if state['mode'] in ['chat', 'chat-instruct'] and not impersonate and not _continue:
            prompt += apply_extensions('bot_prefix', "", state)

        return prompt

    prompt = make_prompt(messages)

    # Handle truncation
    if shared.tokenizer is not None:
        max_length = get_max_prompt_length(state)
        encoded_length = get_encoded_length(prompt)
        while len(messages) > 0 and encoded_length > max_length:

            # Remove old message, save system message
            if len(messages) > 2 and messages[0]['role'] == 'system':
                messages.pop(1)

            # Remove old message when no system message is present
            elif len(messages) > 1 and messages[0]['role'] != 'system':
                messages.pop(0)

            # Resort to truncating the user input
            else:
                user_message = messages[-1]['content']

                # Bisect the truncation point
                left, right = 0, len(user_message)

                while left < right:
                    mid = (left + right + 1) // 2

                    messages[-1]['content'] = user_message[:mid]
                    prompt = make_prompt(messages)
                    encoded_length = get_encoded_length(prompt)

                    if encoded_length <= max_length:
                        left = mid
                    else:
                        right = mid - 1

                messages[-1]['content'] = user_message[:left]
                prompt = make_prompt(messages)
                encoded_length = get_encoded_length(prompt)
                if encoded_length > max_length:
                    logger.error(f"Failed to build the chat prompt. The input is too long for the available context length.\n\nTruncation length: {state['truncation_length']}\nmax_new_tokens: {state['max_new_tokens']} (is it too high?)\nAvailable context length: {max_length}\n")
                    raise ValueError
                else:
                    # Calculate token counts for the log message
                    original_user_tokens = get_encoded_length(user_message)
                    truncated_user_tokens = get_encoded_length(user_message[:left])
                    total_context = max_length + state['max_new_tokens']

                    logger.warning(
                        f"User message truncated from {original_user_tokens} to {truncated_user_tokens} tokens. "
                        f"Context full: {max_length} input tokens ({total_context} total, {state['max_new_tokens']} for output). "
                        f"Increase ctx-size while loading the model to avoid truncation."
                    )

                    break

            prompt = make_prompt(messages)
            encoded_length = get_encoded_length(prompt)

    if also_return_rows:
        return prompt, [message['content'] for message in messages]
    else:
        return prompt


def count_prompt_tokens(text_input, state):
    """Count tokens for current history + input including attachments"""
    if shared.tokenizer is None:
        return "Tokenizer not available"

    try:
        # Handle dict format with text and files
        files = []
        if isinstance(text_input, dict):
            files = text_input.get('files', [])
            text = text_input.get('text', '')
        else:
            text = text_input
            files = []

        # Create temporary history copy to add attachments
        temp_history = copy.deepcopy(state['history'])
        if 'metadata' not in temp_history:
            temp_history['metadata'] = {}

        # Process attachments if any
        if files:
            row_idx = len(temp_history['internal'])
            for file_path in files:
                add_message_attachment(temp_history, row_idx, file_path, is_user=True)

        # Create temp state with modified history
        temp_state = copy.deepcopy(state)
        temp_state['history'] = temp_history

        # Build prompt using existing logic
        prompt = generate_chat_prompt(text, temp_state)
        current_tokens = get_encoded_length(prompt)
        max_tokens = temp_state['truncation_length']

        percentage = (current_tokens / max_tokens) * 100 if max_tokens > 0 else 0

        return f"History + Input:<br/>{current_tokens:,} / {max_tokens:,} tokens ({percentage:.1f}%)"

    except Exception as e:
        logger.error(f"Error counting tokens: {e}")
        return f"Error: {str(e)}"


def get_stopping_strings(state):
    stopping_strings = []
    renderers = []

    if state['mode'] in ['instruct', 'chat-instruct']:
        template = get_compiled_template(state['instruction_template_str'])
        renderer = partial(template.render, add_generation_prompt=False, bos_token=shared.bos_token, eos_token=shared.eos_token)
        renderers.append(renderer)

    if state['mode'] in ['chat']:
        template = get_compiled_template(state['chat_template_str'])
        renderer = partial(template.render, add_generation_prompt=False, name1=state['name1'], name2=state['name2'])
        renderers.append(renderer)

    fake_messages = [
        {"role": "user", "content": "first user message"},
        {"role": "assistant", "content": "first assistant message"},
        {"role": "user", "content": "second user message"},
        {"role": "assistant", "content": "second assistant message"},
    ]

    stopping_strings = []
    for renderer in renderers:
        prompt = renderer(messages=fake_messages)

        # Find positions of each message content
        first_user_end = prompt.find("first user message") + len("first user message")
        first_assistant_start = prompt.find("first assistant message")
        first_assistant_end = prompt.find("first assistant message") + len("first assistant message")
        second_user_start = prompt.find("second user message")
        second_assistant_end = prompt.find("second assistant message") + len("second assistant message")

        # Extract pieces of text potentially containing unique stopping strings
        texts = [
            prompt[first_user_end:first_assistant_start],
            prompt[first_assistant_end:second_user_start],
            prompt[second_assistant_end:]
        ]

        for text in texts:
            stripped_text = text.strip()
            if stripped_text.startswith("<") and ">" in stripped_text:
                stopping_strings.append(stripped_text.split(">")[0] + ">")
            elif stripped_text.startswith("[") and "]" in stripped_text:
                stopping_strings.append(stripped_text.split("]")[0] + "]")
            elif stripped_text.startswith("(") and ")" in stripped_text:
                stopping_strings.append(stripped_text.split(")")[0] + ")")
            elif stripped_text.startswith("{") and "}" in stripped_text:
                stopping_strings.append(stripped_text.split("}")[0] + "}")
            elif ":" in text:
                stopping_strings.append(text.split(":")[0] + ":")

    if 'stopping_strings' in state and isinstance(state['stopping_strings'], list):
        stopping_strings += state.pop('stopping_strings')

    # Remove redundant items that start with another item
    result = [item for item in stopping_strings if not any(item.startswith(other) and item != other for other in stopping_strings)]
    result = list(set(result))

    # Handle GPT-OSS as a special case
    if '<|channel|>final<|message|>' in state['instruction_template_str'] and "<|end|>" in result:
        result.remove("<|end|>")
        result.append("<|result|>")
        result = list(set(result))

    if shared.args.verbose:
        logger.info("STOPPING_STRINGS=")
        pprint.PrettyPrinter(indent=4, sort_dicts=False).pprint(result)
        print()

    return result


def add_message_version(history, role, row_idx, is_current=True):
    key = f"{role}_{row_idx}"
    if 'metadata' not in history:
        history['metadata'] = {}
    if key not in history['metadata']:
        history['metadata'][key] = {}

    if "versions" not in history['metadata'][key]:
        history['metadata'][key]["versions"] = []

    # Determine which index to use for content based on role
    content_idx = 0 if role == 'user' else 1
    current_content = history['internal'][row_idx][content_idx]
    current_visible = history['visible'][row_idx][content_idx]

    history['metadata'][key]["versions"].append({
        "content": current_content,
        "visible_content": current_visible,
        "timestamp": get_current_timestamp()
    })

    if is_current:
        # Set the current_version_index to the newly added version (which is now the last one).
        history['metadata'][key]["current_version_index"] = len(history['metadata'][key]["versions"]) - 1


def add_message_attachment(history, row_idx, file_path, is_user=True):
    """Add a file attachment to a message in history metadata"""
    if 'metadata' not in history:
        history['metadata'] = {}

    key = f"{'user' if is_user else 'assistant'}_{row_idx}"

    if key not in history['metadata']:
        history['metadata'][key] = {"timestamp": get_current_timestamp()}
    if "attachments" not in history['metadata'][key]:
        history['metadata'][key]["attachments"] = []

    # Get file info using pathlib
    path = Path(file_path)
    filename = path.name
    file_extension = path.suffix.lower()

    try:
        # Handle image files
        if file_extension in ['.jpg', '.jpeg', '.png', '.webp', '.bmp', '.gif']:
            # Convert image to base64
            with open(path, 'rb') as f:
                image_data = base64.b64encode(f.read()).decode('utf-8')

            # Determine MIME type from extension
            mime_type_map = {
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.png': 'image/png',
                '.webp': 'image/webp',
                '.bmp': 'image/bmp',
                '.gif': 'image/gif'
            }
            mime_type = mime_type_map.get(file_extension, 'image/jpeg')

            # Format as data URL
            data_url = f"data:{mime_type};base64,{image_data}"

            # Generate unique image ID
            image_id = len([att for att in history['metadata'][key]["attachments"] if att.get("type") == "image"]) + 1

            attachment = {
                "name": filename,
                "type": "image",
                "image_data": data_url,
                "image_id": image_id,
            }
        elif file_extension == '.pdf':
            # Process PDF file
            content = extract_pdf_text(path)
            attachment = {
                "name": filename,
                "type": "application/pdf",
                "content": content,
            }
        elif file_extension == '.docx':
            content = extract_docx_text(path)
            attachment = {
                "name": filename,
                "type": "application/docx",
                "content": content,
            }
        else:
            # Default handling for text files
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()

            attachment = {
                "name": filename,
                "type": "text/plain",
                "content": content,
            }

        history['metadata'][key]["attachments"].append(attachment)
        return attachment  # Return the attachment for reuse
    except Exception as e:
        logger.error(f"Error processing attachment {filename}: {e}")
        return None


def extract_pdf_text(pdf_path):
    """Extract text from a PDF file"""
    import pymupdf

    text = ""
    try:
        with pymupdf.open(pdf_path) as doc:
            for page in doc:
                text += page.get_text() + "\n\n"

        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        return f"[Error extracting PDF text: {str(e)}]"


def extract_docx_text(docx_path):
    """
    Extract text from a .docx file, including headers,
    body (paragraphs and tables), and footers.
    """
    try:
        import docx

        doc = docx.Document(docx_path)
        parts = []

        # 1) Extract non-empty header paragraphs from each section
        for section in doc.sections:
            for para in section.header.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

        # 2) Extract body blocks (paragraphs and tables) in document order
        parent_elm = doc.element.body
        for child in parent_elm.iterchildren():
            if isinstance(child, docx.oxml.text.paragraph.CT_P):
                para = docx.text.paragraph.Paragraph(child, doc)
                text = para.text.strip()
                if text:
                    parts.append(text)

            elif isinstance(child, docx.oxml.table.CT_Tbl):
                table = docx.table.Table(child, doc)
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    parts.append("\t".join(cells))

        # 3) Extract non-empty footer paragraphs from each section
        for section in doc.sections:
            for para in section.footer.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

        return "\n".join(parts)

    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        return f"[Error extracting DOCX text: {str(e)}]"


def generate_search_query(user_message, state):
    """Generate a search query from user message using the LLM"""
    # Augment the user message with search instruction
    augmented_message = f"{user_message}\n\n=====\n\nPlease turn the message above into a short web search query in the same language as the message. Respond with only the search query, nothing else."

    # Use a minimal state for search query generation but keep the full history
    search_state = state.copy()
    search_state['auto_max_new_tokens'] = True
    search_state['enable_thinking'] = False
    search_state['reasoning_effort'] = 'low'
    search_state['start_with'] = ""

    # Generate the full prompt using existing history + augmented message
    formatted_prompt = generate_chat_prompt(augmented_message, search_state)

    query = ""
    for reply in generate_reply(formatted_prompt, search_state, stopping_strings=[], is_chat=True):
        query = reply

    # Check for thinking block delimiters and extract content after them
    if "</think>" in query:
        query = query.rsplit("</think>", 1)[1]
    elif "<|start|>assistant<|channel|>final<|message|>" in query:
        query = query.rsplit("<|start|>assistant<|channel|>final<|message|>", 1)[1]
    elif "<|channel|>final<|message|>" in query:
        query = query.rsplit("<|channel|>final<|message|>", 1)[1]
    elif "</seed:think>" in query:
        query = query.rsplit("</seed:think>", 1)[1]

    # Strip and remove surrounding quotes if present
    query = query.strip()
    if len(query) >= 2 and query.startswith('"') and query.endswith('"'):
        query = query[1:-1]

    return query


def chatbot_wrapper(text, state, regenerate=False, _continue=False, loading_message=True, for_ui=False):
    # Handle dict format with text and files
    files = []
    if isinstance(text, dict):
        files = text.get('files', [])
        text = text.get('text', '')

    history = state['history']
    output = copy.deepcopy(history)
    output = apply_extensions('history', output)
    state = apply_extensions('state', state)

    # Handle GPT-OSS as a special case
    if '<|channel|>final<|message|>' in state['instruction_template_str']:
        state['skip_special_tokens'] = False

    # Let the jinja2 template handle the BOS token
    if state['mode'] in ['instruct', 'chat-instruct']:
        state['add_bos_token'] = False

    # Initialize metadata if not present
    if 'metadata' not in output:
        output['metadata'] = {}

    visible_text = None
    stopping_strings = get_stopping_strings(state)
    is_stream = state['stream']

    # Prepare the input
    if not (regenerate or _continue):
        visible_text = html.escape(text)

        # Process file attachments and store in metadata
        row_idx = len(output['internal'])

        # Add attachments to metadata only, not modifying the message text
        for file_path in files:
            add_message_attachment(output, row_idx, file_path, is_user=True)

        # Add web search results as attachments if enabled
        if state.get('enable_web_search', False):
            search_query = generate_search_query(text, state)
            add_web_search_attachments(output, row_idx, text, search_query, state)

        # Apply extensions
        text, visible_text = apply_extensions('chat_input', text, visible_text, state)
        text = apply_extensions('input', text, state, is_chat=True)

        # Current row index
        output['internal'].append([text, ''])
        output['visible'].append([visible_text, ''])
        # Add metadata with timestamp
        update_message_metadata(output['metadata'], "user", row_idx, timestamp=get_current_timestamp())

        # *Is typing...*
        if loading_message:
            yield {
                'visible': output['visible'][:-1] + [[output['visible'][-1][0], shared.processing_message]],
                'internal': output['internal'],
                'metadata': output['metadata']
            }
    else:
        text, visible_text = output['internal'][-1][0], output['visible'][-1][0]
        if regenerate and not state.get('_tool_turn'):
            row_idx = len(output['internal']) - 1

            # Store the old response as a version before regenerating
            if not output['metadata'].get(f"assistant_{row_idx}", {}).get('versions'):
                add_message_version(output, "assistant", row_idx, is_current=False)

            # Add new empty version (will be filled during streaming)
            key = f"assistant_{row_idx}"
            output['metadata'][key]["versions"].append({
                "content": "",
                "visible_content": "",
                "timestamp": get_current_timestamp()
            })
            output['metadata'][key]["current_version_index"] = len(output['metadata'][key]["versions"]) - 1

            if loading_message:
                yield {
                    'visible': output['visible'][:-1] + [[visible_text, shared.processing_message]],
                    'internal': output['internal'][:-1] + [[text, '']],
                    'metadata': output['metadata']
                }
        elif _continue:
            last_reply = [output['internal'][-1][1], output['visible'][-1][1]]
            if loading_message:
                yield {
                    'visible': output['visible'][:-1] + [[visible_text, last_reply[1] + '...']],
                    'internal': output['internal'],
                    'metadata': output['metadata']
                }

    row_idx = len(output['internal']) - 1

    # Collect image attachments for multimodal generation from the entire history
    all_image_attachments = []
    if 'metadata' in output:
        for i in range(len(output['internal'])):
            user_key = f"user_{i}"
            if user_key in output['metadata'] and "attachments" in output['metadata'][user_key]:
                for attachment in output['metadata'][user_key]["attachments"]:
                    if attachment.get("type") == "image":
                        all_image_attachments.append(attachment)

    # Add all collected image attachments to state for the generation
    if all_image_attachments:
        state['image_attachments'] = all_image_attachments

    # Generate the prompt
    kwargs = {
        '_continue': _continue,
        'history': output if _continue else {
            k: (v[:-1] if k in ['internal', 'visible'] else v)
            for k, v in output.items()
        }
    }

    prompt = apply_extensions('custom_generate_chat_prompt', text, state, **kwargs)
    if prompt is None:
        prompt = generate_chat_prompt(text, state, **kwargs)

    # Add timestamp for assistant's response at the start of generation
    update_message_metadata(output['metadata'], "assistant", row_idx, timestamp=get_current_timestamp(), model_name=shared.model_name)

    # Detect if the template appended a thinking start tag to the prompt
    thinking_prefix = None
    if not _continue:
        stripped_prompt = prompt.rstrip('\n')
        for start_tag, end_tag, content_tag in THINKING_FORMATS:
            if start_tag is not None and stripped_prompt.endswith(start_tag):
                thinking_prefix = start_tag
                break

    # When tools are active, buffer streaming output during potential tool
    # call generation to prevent raw markup from leaking into the display.
    _check_tool_markers = bool(state.get('tools'))
    _last_visible_before_tool_buffer = None
    if _check_tool_markers:
        from modules.tool_parsing import streaming_tool_buffer_check, detect_tool_call_format
        _tool_names = [t['function']['name'] for t in state['tools'] if 'function' in t and 'name' in t['function']]
        _template_str = state.get('instruction_template_str', '') if state.get('mode') == 'instruct' else state.get('chat_template_str', '')
        _, _streaming_markers, _check_bare_names = detect_tool_call_format(_template_str)

    # Generate
    reply = None
    for j, reply in enumerate(generate_reply(prompt, state, stopping_strings=stopping_strings, is_chat=True, for_ui=for_ui)):

        # Prepend thinking tag if the template appended it to the prompt
        if thinking_prefix:
            reply = thinking_prefix + reply

        # Extract the reply
        if state['mode'] in ['chat', 'chat-instruct']:
            if not _continue:
                reply = reply.lstrip()

            if reply.startswith(state['name2'] + ':'):
                reply = reply[len(state['name2'] + ':'):]
            elif reply.startswith(state['name1'] + ':'):
                reply = reply[len(state['name1'] + ':'):]

            visible_reply = re.sub("(<USER>|<user>|{{user}})", state['name1'], reply)
        else:
            visible_reply = reply

        visible_reply = html.escape(visible_reply)

        if shared.stop_everything:
            if not state.get('_skip_output_extensions'):
                output['visible'][-1][1] = apply_extensions('output', output['visible'][-1][1], state, is_chat=True)

            yield output
            return

        if _continue:
            output['internal'][-1] = [text, last_reply[0] + reply]
            output['visible'][-1] = [visible_text, last_reply[1] + visible_reply]
        elif not (j == 0 and visible_reply.strip() == ''):
            output['internal'][-1] = [text, reply.lstrip(' ')]
            output['visible'][-1] = [visible_text, visible_reply.lstrip(' ')]

        # Keep version metadata in sync during streaming (for regeneration)
        if regenerate and not state.get('_tool_turn'):
            row_idx = len(output['internal']) - 1
            key = f"assistant_{row_idx}"
            current_idx = output['metadata'][key]['current_version_index']
            output['metadata'][key]['versions'][current_idx].update({
                'content': output['internal'][row_idx][1],
                'visible_content': output['visible'][row_idx][1]
            })

        if is_stream:
            if _check_tool_markers:
                if streaming_tool_buffer_check(output['internal'][-1][1], markers=_streaming_markers, tool_names=_tool_names, check_bare_names=_check_bare_names):
                    continue
                _last_visible_before_tool_buffer = output['visible'][-1][1]

            yield output

    if _continue:
        # Reprocess the entire internal text for extensions (like translation).
        # Skip entirely when the visible text contains <tool_call> markers,
        # since those only exist in visible (internal is cleared after each tool
        # execution) and rebuilding from internal would destroy them. Output
        # extensions also can't handle the raw <tool_call> markup safely.
        if '<tool_call>' not in output['visible'][-1][1]:
            full_internal = output['internal'][-1][1]
            if state['mode'] in ['chat', 'chat-instruct']:
                full_visible = re.sub("(<USER>|<user>|{{user}})", state['name1'], full_internal)
            else:
                full_visible = full_internal

            full_visible = html.escape(full_visible)
            if not state.get('_skip_output_extensions'):
                output['visible'][-1][1] = apply_extensions('output', full_visible, state, is_chat=True)
    else:
        if not state.get('_skip_output_extensions'):
            output['visible'][-1][1] = apply_extensions('output', output['visible'][-1][1], state, is_chat=True)

    # Final sync for version metadata (in case streaming was disabled)
    if regenerate and not state.get('_tool_turn'):
        row_idx = len(output['internal']) - 1
        key = f"assistant_{row_idx}"
        current_idx = output['metadata'][key]['current_version_index']
        output['metadata'][key]['versions'][current_idx].update({
            'content': output['internal'][row_idx][1],
            'visible_content': output['visible'][row_idx][1]
        })

    # When tool markers were detected during streaming, restore the last
    # visible text from before buffering started so raw markup doesn't flash
    # in the UI.  The internal text is left intact so the caller can still
    # parse tool calls from it.
    if is_stream and _check_tool_markers and streaming_tool_buffer_check(output['internal'][-1][1], markers=_streaming_markers, tool_names=_tool_names, check_bare_names=_check_bare_names):
        output['visible'][-1][1] = _last_visible_before_tool_buffer or ''

    yield output


def impersonate_wrapper(textbox, state):
    text = textbox['text']
    static_output = chat_html_wrapper(state['history'], state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'])

    prompt = generate_chat_prompt('', state, impersonate=True)
    stopping_strings = get_stopping_strings(state)

    textbox['text'] = text + '...'
    yield textbox, static_output
    reply = None
    for reply in generate_reply(prompt + text, state, stopping_strings=stopping_strings, is_chat=True):
        textbox['text'] = (text + reply).lstrip(' ')
        yield textbox, static_output
        if shared.stop_everything:
            return


def generate_chat_reply(text, state, regenerate=False, _continue=False, loading_message=True, for_ui=False):
    history = state['history']
    if regenerate or _continue:
        text = ''
        if (len(history['visible']) == 1 and not history['visible'][0][0]) or len(history['internal']) == 0:
            yield history
            return

    for history in chatbot_wrapper(text, state, regenerate=regenerate, _continue=_continue, loading_message=loading_message, for_ui=for_ui):
        yield history


def character_is_loaded(state, raise_exception=False):
    if state['mode'] in ['chat', 'chat-instruct'] and state['name2'] == '':
        logger.error('It looks like no character is loaded. Please load one under Parameters > Character.')
        if raise_exception:
            raise ValueError

        return False
    else:
        return True


def generate_chat_reply_wrapper(text, state, regenerate=False, _continue=False):
    '''
    Same as above but returns HTML for the UI.
    When tools are selected, wraps generation in a loop that detects
    tool calls, executes them, and re-generates until the model stops.
    All tool output is consolidated into a single visible chat bubble
    using metadata['assistant_N']['tool_sequence'].
    '''

    if not character_is_loaded(state):
        return

    if state['start_with'] != '' and not _continue:
        if regenerate:
            text, state['history'] = remove_last_message(state['history'])
            regenerate = False

        _continue = True
        send_dummy_message(text, state)
        send_dummy_reply(state['start_with'], state)

    # On regenerate, clear old tool_sequence metadata so it gets rebuilt.
    # Save it first so it can be stored per-version below.
    # This must happen after the start_with logic above, which may remove
    # and re-add messages, changing which row we operate on.
    _old_tool_sequence = None
    if regenerate:
        history = state['history']
        meta = history.get('metadata', {})
        row_idx = len(history['internal']) - 1
        if row_idx >= 0:
            _old_tool_sequence = meta.get(f'assistant_{row_idx}', {}).pop('tool_sequence', None)

    # Load tools if any are selected
    selected = state.get('selected_tools', [])
    parse_tool_call = None
    _tool_parsers = None
    if selected:
        from modules.tool_use import load_tools, execute_tool
        from modules.tool_parsing import parse_tool_call, get_tool_call_id, detect_tool_call_format

    if selected:
        tool_defs, tool_executors = load_tools(selected)
        state['tools'] = tool_defs
        tool_func_names = [t['function']['name'] for t in tool_defs]
        _template_str = state.get('instruction_template_str', '') if state.get('mode') == 'instruct' else state.get('chat_template_str', '')
        _tool_parsers, _, _ = detect_tool_call_format(_template_str)
    else:
        tool_func_names = None

    visible_prefix = []  # Accumulated tool call summaries + results
    last_save_time = time.monotonic()
    save_interval = 8
    _tool_turn = 0
    while True:
        history = state['history']

        # Turn 0: use original flags; turns 2+: regenerate into the same entry.
        # _tool_turn tells chatbot_wrapper to skip version creation/sync so
        # that intermediate tool-loop regenerations don't pollute swipe history.
        if _tool_turn > 0:
            state['_tool_turn'] = True
            state['_skip_output_extensions'] = True

        regen = regenerate if _tool_turn == 0 else True
        cont = _continue if _tool_turn == 0 else False
        cur_text = text if _tool_turn == 0 else ''

        for i, history in enumerate(generate_chat_reply(cur_text, state, regen, cont, loading_message=True, for_ui=True)):
            # Prepend accumulated tool output to visible reply for display.
            # Save and restore the original to prevent the markers from leaking
            # back into chatbot_wrapper's shared output object, which would cause
            # duplication on the next yield.
            _original_visible = history['visible'][-1][1] if visible_prefix else None
            if visible_prefix:
                history['visible'][-1][1] = '\n\n'.join(visible_prefix + [_original_visible])

            yield chat_html_wrapper(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'], last_message_only=(i > 0)), history

            if visible_prefix:
                history['visible'][-1][1] = _original_visible

            if i == 0:
                # Save old tool_sequence into version 0 (created by chatbot_wrapper
                # on the first yield).  Only needed on the first regeneration when
                # versions didn't previously exist.
                if _old_tool_sequence is not None and _tool_turn == 0:
                    _ri = len(history['internal']) - 1
                    _versions = history.get('metadata', {}).get(f'assistant_{_ri}', {}).get('versions', [])
                    if _versions and 'tool_sequence' not in _versions[0]:
                        _versions[0]['tool_sequence'] = _old_tool_sequence
                    _old_tool_sequence = None

                time.sleep(0.125)

            current_time = time.monotonic()
            if i == 0 or (current_time - last_save_time) >= save_interval:
                save_history(history, state['unique_id'], state['character_menu'], state['mode'])
                last_save_time = current_time

            # Early stop on tool call detection
            if tool_func_names and parse_tool_call(history['internal'][-1][1], tool_func_names, parsers=_tool_parsers):
                break

        # Save the model's visible output before re-applying visible_prefix,
        # so we can extract thinking content from just this turn's output.
        _model_visible = history['visible'][-1][1]

        # Recover visible_prefix from existing visible text (e.g. on Continue
        # after a previous session had tool calls). Extract all <tool_call>
        # blocks and any text between them (thinking blocks, intermediate text).
        if tool_func_names and not visible_prefix and _model_visible:
            tc_matches = list(re.finditer(r'<tool_call>.*?</tool_call>', _model_visible, re.DOTALL))
            if tc_matches:
                prefix_end = tc_matches[-1].end()
                prefix = _model_visible[:prefix_end].strip()
                if prefix:
                    visible_prefix = [prefix]
                _model_visible = _model_visible[prefix_end:].strip()

        # Re-apply visible prefix to the final state after streaming completes.
        # This is safe because we're no longer sharing the object with chatbot_wrapper.
        if visible_prefix:
            history['visible'][-1][1] = '\n\n'.join(visible_prefix + [_model_visible])

        if tool_func_names:
            save_history(history, state['unique_id'], state['character_menu'], state['mode'])

        # Check for tool calls
        if not tool_func_names or shared.stop_everything:
            break

        answer = history['internal'][-1][1]
        parsed_calls, content_prefix = parse_tool_call(answer, tool_func_names, return_prefix=True, parsers=_tool_parsers) if answer else (None, '')

        if not parsed_calls:
            break  # No tool calls — done

        # --- Process tool calls ---
        row_idx = len(history['internal']) - 1
        meta = history.get('metadata', {})
        seq = meta.setdefault(f'assistant_{row_idx}', {}).setdefault('tool_sequence', [])

        def _render():
            return chat_html_wrapper(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'])

        # Serialize tool calls and build display headers in one pass
        serialized = []
        tc_headers = []
        for tc in parsed_calls:
            tc['id'] = get_tool_call_id()
            fn_name = tc['function']['name']
            fn_args = tc['function'].get('arguments', {})

            serialized.append({
                'id': tc['id'],
                'type': 'function',
                'function': {
                    'name': fn_name,
                    'arguments': json.dumps(fn_args) if isinstance(fn_args, dict) else fn_args
                }
            })

            if isinstance(fn_args, dict) and fn_args:
                args_summary = ', '.join(f'{k}={json.dumps(v, ensure_ascii=False)}' for k, v in fn_args.items())
            elif isinstance(fn_args, dict):
                args_summary = ''
            else:
                args_summary = str(fn_args)

            tc_headers.append(f'{fn_name}({args_summary})')

        seq_entry = {'tool_calls': serialized}
        if content_prefix.strip():
            # Strip GPT-OSS channel tokens so they don't get double-wrapped
            # by the template (which adds its own channel markup).
            clean = content_prefix.strip()
            if '<|channel|>' in clean and '<|message|>' in clean:
                inner = clean.split('<|message|>', 1)[1]
                if '<|end|>' in inner:
                    inner = inner.split('<|end|>', 1)[0]
                clean = inner.strip()
            if clean:
                seq_entry['content'] = clean
        seq.append(seq_entry)

        # Clear internal (raw tool markup)
        history['internal'][-1][1] = ''

        # Preserve thinking block and intermediate text from this turn.
        # content_prefix is the raw text before tool call syntax (returned
        # by parse_tool_call); HTML-escape it and extract thinking to get
        # the content the user should see.
        content_text = html.escape(content_prefix)
        thinking_content, intermediate = extract_thinking_block(content_text)
        if thinking_content:
            visible_prefix.append(f'&lt;think&gt;\n{thinking_content}\n&lt;/think&gt;')
        if intermediate and intermediate.strip():
            visible_prefix.append(intermediate.strip())

        # Show placeholder accordions with "..." before execution starts
        # (tool calls may be slow, e.g. web search).
        pending_placeholders = [f'<tool_call>{h}\n...\n</tool_call>' for h in tc_headers]
        history['visible'][-1][1] = '\n\n'.join(visible_prefix + pending_placeholders)
        yield _render(), history

        # Execute tools, store results, and replace placeholders with real results
        for i, tc in enumerate(parsed_calls):
            # Check for stop request before each tool execution
            if shared.stop_everything:
                for j in range(i, len(parsed_calls)):
                    seq.append({'role': 'tool', 'content': 'Tool execution was cancelled by the user.', 'tool_call_id': parsed_calls[j]['id']})
                    pending_placeholders[j] = f'<tool_call>{tc_headers[j]}\nCancelled\n</tool_call>'

                history['visible'][-1][1] = '\n\n'.join(visible_prefix + pending_placeholders)
                yield _render(), history
                break

            fn_name = tc['function']['name']
            fn_args = tc['function'].get('arguments', {})
            result = execute_tool(fn_name, fn_args, tool_executors)

            seq.append({'role': 'tool', 'content': result, 'tool_call_id': tc['id']})
            try:
                pretty_result = json.dumps(json.loads(result), indent=2, ensure_ascii=False)
            except (json.JSONDecodeError, TypeError):
                pretty_result = result

            # Replace the placeholder with the real result
            pending_placeholders[i] = f'<tool_call>{tc_headers[i]}\n{pretty_result}\n</tool_call>'
            history['visible'][-1][1] = '\n\n'.join(visible_prefix + pending_placeholders)
            yield _render(), history

        # Move completed tool calls into visible_prefix for next turns
        visible_prefix.extend(pending_placeholders)
        history['visible'][-1][1] = '\n\n'.join(visible_prefix)
        save_history(history, state['unique_id'], state['character_menu'], state['mode'])

        state['history'] = history
        _tool_turn += 1

    state.pop('_tool_turn', None)

    # If output extensions were deferred during tool turns, apply them now
    # to the final model response only (not to tool call markers).
    if state.pop('_skip_output_extensions', None):
        _model_visible = apply_extensions('output', _model_visible, state, is_chat=True)
        if visible_prefix:
            history['visible'][-1][1] = '\n\n'.join(visible_prefix + [_model_visible])
        else:
            history['visible'][-1][1] = _model_visible

        yield chat_html_wrapper(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu']), history

    state['history'] = history

    # Sync version metadata so swipes show the full visible (with tool prefix)
    if visible_prefix and history.get('metadata'):
        row_idx = len(history['internal']) - 1
        key = f"assistant_{row_idx}"
        meta_entry = history['metadata'].get(key, {})
        if 'versions' in meta_entry and 'current_version_index' in meta_entry:
            current_idx = meta_entry['current_version_index']
            if current_idx < len(meta_entry['versions']):
                version_update = {
                    'content': history['internal'][row_idx][1],
                    'visible_content': history['visible'][row_idx][1]
                }
                ts = meta_entry.get('tool_sequence')
                if ts is not None:
                    version_update['tool_sequence'] = ts
                meta_entry['versions'][current_idx].update(version_update)

    save_history(history, state['unique_id'], state['character_menu'], state['mode'])


def remove_last_message(history):
    if 'metadata' not in history:
        history['metadata'] = {}

    if len(history['visible']) > 0 and history['internal'][-1][0] != '<|BEGIN-VISIBLE-CHAT|>':
        row_idx = len(history['internal']) - 1
        last = history['visible'].pop()
        history['internal'].pop()

        # Remove metadata directly by known keys
        if f"user_{row_idx}" in history['metadata']:
            del history['metadata'][f"user_{row_idx}"]
        if f"assistant_{row_idx}" in history['metadata']:
            del history['metadata'][f"assistant_{row_idx}"]
    else:
        last = ['', '']

    return html.unescape(last[0]), history


def send_dummy_message(text, state):
    history = state['history']

    # Handle both dict and string inputs
    if isinstance(text, dict):
        text = text['text']

    # Initialize metadata if not present
    if 'metadata' not in history:
        history['metadata'] = {}

    row_idx = len(history['internal'])
    history['visible'].append([html.escape(text), ''])
    history['internal'].append([apply_extensions('input', text, state, is_chat=True), ''])
    update_message_metadata(history['metadata'], "user", row_idx, timestamp=get_current_timestamp())

    return history


def send_dummy_reply(text, state):
    history = state['history']

    # Handle both dict and string inputs
    if isinstance(text, dict):
        text = text['text']

    # Initialize metadata if not present
    if 'metadata' not in history:
        history['metadata'] = {}

    if len(history['visible']) > 0 and not history['visible'][-1][1] == '':
        row_idx = len(history['internal'])
        history['visible'].append(['', ''])
        history['internal'].append(['', ''])
        # We don't need to add system metadata

    row_idx = len(history['internal']) - 1
    history['visible'][-1][1] = html.escape(text)
    history['internal'][-1][1] = apply_extensions('input', text, state, is_chat=True)
    update_message_metadata(history['metadata'], "assistant", row_idx, timestamp=get_current_timestamp())

    return history


def redraw_html(history, name1, name2, mode, style, character, reset_cache=False):
    return chat_html_wrapper(history, name1, name2, mode, style, character, reset_cache=reset_cache)


def start_new_chat(state, unique_id=None):
    mode = state['mode']
    # Initialize with empty metadata dictionary
    history = {'internal': [], 'visible': [], 'metadata': {}}

    if mode != 'instruct':
        greeting = replace_character_names(state['greeting'], state['name1'], state['name2'])
        if greeting != '':
            history['internal'] += [['<|BEGIN-VISIBLE-CHAT|>', greeting]]
            history['visible'] += [['', apply_extensions('output', html.escape(greeting), state, is_chat=True)]]

            # Add timestamp for assistant's greeting
            update_message_metadata(history['metadata'], "assistant", 0, timestamp=get_current_timestamp())

    if unique_id is None:
        unique_id = datetime.now().strftime('%Y%m%d-%H-%M-%S')

    save_history(history, unique_id, state['character_menu'], state['mode'])

    return history


def get_history_file_path(unique_id, character, mode):
    if mode == 'instruct':
        p = shared.user_data_dir / 'logs' / 'instruct' / f'{unique_id}.json'
    else:
        p = shared.user_data_dir / 'logs' / 'chat' / character / f'{unique_id}.json'

    return p


def save_history(history, unique_id, character, mode):
    if shared.args.multi_user:
        return

    if unique_id and unique_id.startswith('incognito-'):
        return

    p = get_history_file_path(unique_id, character, mode)
    if not p.parent.is_dir():
        p.parent.mkdir(parents=True)

    with _history_file_lock:
        with open(p, 'w', encoding='utf-8') as f:
            f.write(json.dumps(history, indent=4, ensure_ascii=False))


def rename_history(old_id, new_id, character, mode):
    if shared.args.multi_user:
        return

    old_p = get_history_file_path(old_id, character, mode)
    new_p = get_history_file_path(new_id, character, mode)
    if new_p.parent != old_p.parent:
        logger.error(f"The following path is not allowed: \"{new_p}\".")
    elif new_p == old_p:
        logger.info("The provided path is identical to the old one.")
    elif new_p.exists():
        logger.error(f"The new path already exists and will not be overwritten: \"{new_p}\".")
    else:
        logger.info(f"Renaming \"{old_p}\" to \"{new_p}\"")
        old_p.rename(new_p)


def get_paths(state):
    if state['mode'] == 'instruct':
        return (shared.user_data_dir / 'logs' / 'instruct').glob('*.json')
    else:
        character = state['character_menu']

        # Handle obsolete filenames and paths
        old_p = shared.user_data_dir / 'logs' / f'{character}_persistent.json'
        new_p = shared.user_data_dir / 'logs' / f'persistent_{character}.json'
        if old_p.exists():
            logger.warning(f"Renaming \"{old_p}\" to \"{new_p}\"")
            old_p.rename(new_p)

        if new_p.exists():
            unique_id = datetime.now().strftime('%Y%m%d-%H-%M-%S')
            p = get_history_file_path(unique_id, character, state['mode'])
            logger.warning(f"Moving \"{new_p}\" to \"{p}\"")
            p.parent.mkdir(exist_ok=True)
            new_p.rename(p)

        return (shared.user_data_dir / 'logs' / 'chat' / character).glob('*.json')


def find_all_histories(state):
    if shared.args.multi_user:
        return ['']

    paths = get_paths(state)
    histories = sorted(paths, key=lambda x: x.stat().st_mtime, reverse=True)
    return [path.stem for path in histories]


def find_all_histories_with_first_prompts(state):
    if shared.args.multi_user:
        return []

    paths = get_paths(state)
    histories = sorted(paths, key=lambda x: x.stat().st_mtime, reverse=True)

    result = []
    for i, path in enumerate(histories):
        filename = path.stem
        file_content = ""
        with open(path, 'r', encoding='utf-8') as f:
            file_content = f.read()

        if state['search_chat'] and state['search_chat'] not in file_content:
            continue

        data = json.loads(file_content)
        if re.match(r'^[0-9]{8}-[0-9]{2}-[0-9]{2}-[0-9]{2}$', filename):
            first_prompt = ""
            if data and 'visible' in data and len(data['visible']) > 0:
                if len(data['internal']) > 0 and data['internal'][0][0] == '<|BEGIN-VISIBLE-CHAT|>':
                    if len(data['visible']) > 1:
                        first_prompt = html.unescape(data['visible'][1][0])
                    elif i == 0:
                        first_prompt = "New chat"
                else:
                    first_prompt = html.unescape(data['visible'][0][0])
            elif i == 0:
                first_prompt = "New chat"
        else:
            first_prompt = filename

        first_prompt = first_prompt.strip()

        # Truncate the first prompt if it's longer than 30 characters
        if len(first_prompt) > 30:
            first_prompt = first_prompt[:30 - 3] + '...'

        result.append((first_prompt, filename))

    return result


def load_latest_history(state):
    '''
    Loads the latest history for the given character in chat or chat-instruct
    mode, or the latest instruct history for instruct mode.
    '''

    if shared.args.multi_user:
        return start_new_chat(state), None

    histories = find_all_histories(state)

    if len(histories) > 0:
        # Try to load the last visited chat for this character/mode
        chat_state = load_last_chat_state()
        key = get_chat_state_key(state['character_menu'], state['mode'])
        last_chat_id = chat_state.get("last_chats", {}).get(key)

        # If we have a stored last chat and it still exists, use it
        if last_chat_id and last_chat_id in histories:
            unique_id = last_chat_id
        else:
            # Fall back to most recent (current behavior)
            unique_id = histories[0]

        history = load_history(unique_id, state['character_menu'], state['mode'])
        return history, unique_id
    else:
        return start_new_chat(state), None


def load_history_after_deletion(state, idx):
    '''
    Loads the latest history for the given character in chat or chat-instruct
    mode, or the latest instruct history for instruct mode.
    '''
    import gradio as gr

    if shared.args.multi_user:
        return start_new_chat(state)

    histories = find_all_histories_with_first_prompts(state)
    idx = min(int(idx), len(histories) - 1)
    idx = max(0, idx)

    if len(histories) > 0:
        history = load_history(histories[idx][1], state['character_menu'], state['mode'])
    else:
        history = start_new_chat(state)
        histories = find_all_histories_with_first_prompts(state)

    return history, gr.update(choices=histories, value=histories[idx][1])


def update_character_menu_after_deletion(idx):
    import gradio as gr
    characters = utils.get_available_characters()
    idx = min(int(idx), len(characters) - 1)
    idx = max(0, idx)
    return gr.update(choices=characters, value=characters[idx])


def get_chat_state_key(character, mode):
    """Generate a key for storing last chat state"""
    if mode == 'instruct':
        return 'instruct'
    else:
        return f"chat_{character}"


def load_last_chat_state():
    """Load the last chat state from file"""
    state_file = shared.user_data_dir / 'logs' / 'chat_state.json'
    if state_file.exists():
        try:
            with open(state_file, 'r', encoding='utf-8') as f:
                return json.loads(f.read())
        except Exception:
            pass

    return {"last_chats": {}}


def save_last_chat_state(character, mode, unique_id):
    """Save the last visited chat for a character/mode"""
    if shared.args.multi_user:
        return

    if unique_id and unique_id.startswith('incognito-'):
        return

    state = load_last_chat_state()
    key = get_chat_state_key(character, mode)
    state["last_chats"][key] = unique_id

    state_file = shared.user_data_dir / 'logs' / 'chat_state.json'
    state_file.parent.mkdir(exist_ok=True)
    with open(state_file, 'w', encoding='utf-8') as f:
        f.write(json.dumps(state, indent=2))


def load_history(unique_id, character, mode):
    p = get_history_file_path(unique_id, character, mode)

    if not p.exists():
        return {'internal': [], 'visible': [], 'metadata': {}}

    f = json.loads(open(p, 'rb').read())
    if 'internal' in f and 'visible' in f:
        history = f
    else:
        history = {
            'internal': f['data'],
            'visible': f['data_visible']
        }

    # Add metadata if it doesn't exist
    if 'metadata' not in history:
        history['metadata'] = {}
        # Add placeholder timestamps for existing messages
        for i, (user_msg, asst_msg) in enumerate(history['internal']):
            if user_msg and user_msg != '<|BEGIN-VISIBLE-CHAT|>':
                update_message_metadata(history['metadata'], "user", i, timestamp="")
            if asst_msg:
                update_message_metadata(history['metadata'], "assistant", i, timestamp="")

    return history


def load_history_json(file, history):
    try:
        file = file.decode('utf-8')
        f = json.loads(file)
        if 'internal' in f and 'visible' in f:
            history = f
        else:
            history = {
                'internal': f['data'],
                'visible': f['data_visible']
            }

        # Add metadata if it doesn't exist
        if 'metadata' not in history:
            history['metadata'] = {}
            # Add placeholder timestamps
            for i, (user_msg, asst_msg) in enumerate(history['internal']):
                if user_msg and user_msg != '<|BEGIN-VISIBLE-CHAT|>':
                    update_message_metadata(history['metadata'], "user", i, timestamp="")
                if asst_msg:
                    update_message_metadata(history['metadata'], "assistant", i, timestamp="")

        return history
    except Exception:
        return history


def delete_history(unique_id, character, mode):
    p = get_history_file_path(unique_id, character, mode)
    delete_file(p)


def replace_character_names(text, name1, name2):
    text = text.replace('{{user}}', name1).replace('{{char}}', name2)
    return text.replace('<USER>', name1).replace('<BOT>', name2)


def generate_pfp_cache(character):
    cache_folder = Path(shared.args.disk_cache_dir)
    if not cache_folder.exists():
        cache_folder.mkdir()

    for path in [shared.user_data_dir / 'characters' / f"{character}.{extension}" for extension in ['png', 'jpg', 'jpeg']]:
        if path.exists():
            original_img = Image.open(path)
            # Define file paths
            pfp_path = Path(f'{cache_folder}/pfp_character.png')
            thumb_path = Path(f'{cache_folder}/pfp_character_thumb.png')

            # Save main picture and thumbnail
            original_img.save(pfp_path, format='PNG')
            thumb = make_thumbnail(original_img)
            thumb.save(thumb_path, format='PNG')

            # Return the path to the thumbnail, not the in-memory PIL Image object.
            return str(thumb_path)

    return None


def load_character(character, name1, name2):
    context = greeting = ""
    greeting_field = 'greeting'
    picture = None

    filepath = None
    for extension in ["yml", "yaml", "json"]:
        filepath = shared.user_data_dir / 'characters' / f'{character}.{extension}'
        if filepath.exists():
            break

    if filepath is None or not filepath.exists():
        logger.error(f"Could not find the character \"{character}\" inside {shared.user_data_dir}/characters. No character has been loaded.")
        raise ValueError

    file_contents = open(filepath, 'r', encoding='utf-8').read()
    data = json.loads(file_contents) if extension == "json" else yaml.safe_load(file_contents)
    cache_folder = Path(shared.args.disk_cache_dir)

    for path in [Path(f"{cache_folder}/pfp_character.png"), Path(f"{cache_folder}/pfp_character_thumb.png")]:
        if path.exists():
            path.unlink()

    picture = generate_pfp_cache(character)

    # Finding the bot's name
    for k in ['name', 'bot', '<|bot|>', 'char_name']:
        if k in data and data[k] != '':
            name2 = data[k]
            break

    # Find the user name (if any)
    for k in ['your_name', 'user', '<|user|>']:
        if k in data and data[k] != '':
            name1 = data[k]
            break

    if 'context' in data:
        context = data['context'].strip()
    elif "char_persona" in data:
        context = build_pygmalion_style_context(data)
        greeting_field = 'char_greeting'

    greeting = data.get(greeting_field, greeting)
    return name1, name2, picture, greeting, context


def restore_character_for_ui(state):
    """Reset character fields to the currently loaded character's saved values"""
    if state['character_menu'] and state['character_menu'] != 'None':
        try:
            name1, name2, picture, greeting, context = load_character(state['character_menu'], state['name1'], state['name2'])

            state['name2'] = name2
            state['greeting'] = greeting
            state['context'] = context
            state['character_picture'] = picture  # This triggers cache update via generate_pfp_cache

            return state, name2, context, greeting, picture

        except Exception as e:
            logger.error(f"Failed to reset character '{state['character_menu']}': {e}")
            return clear_character_for_ui(state)
    else:
        return clear_character_for_ui(state)


def clear_character_for_ui(state):
    """Clear all character fields and picture cache"""
    state['name2'] = shared.settings['name2']
    state['context'] = shared.settings['context']
    state['greeting'] = shared.settings['greeting']
    state['character_picture'] = None

    # Clear the cache files
    cache_folder = Path(shared.args.disk_cache_dir)
    for cache_file in ['pfp_character.png', 'pfp_character_thumb.png']:
        cache_path = Path(f'{cache_folder}/{cache_file}')
        if cache_path.exists():
            cache_path.unlink()

    return state, state['name2'], state['context'], state['greeting'], None


@functools.cache
def load_character_memoized(character, name1, name2):
    return load_character(character, name1, name2)


@functools.cache
def load_instruction_template_memoized(template):
    from modules.models_settings import load_instruction_template
    return load_instruction_template(template)


def upload_character(file, img_path, tavern=False):
    import gradio as gr
    img = open_image_safely(img_path)
    decoded_file = file if isinstance(file, str) else file.decode('utf-8')
    try:
        data = json.loads(decoded_file)
    except Exception:
        data = yaml.safe_load(decoded_file)

    if 'char_name' in data:
        name = sanitize_filename(data['char_name'])
        greeting = data['char_greeting']
        context = build_pygmalion_style_context(data)
        yaml_data = generate_character_yaml(name, greeting, context)
    else:
        name = sanitize_filename(data['name'])
        yaml_data = generate_character_yaml(data['name'], data['greeting'], data['context'])

    outfile_name = name
    i = 1
    while (shared.user_data_dir / 'characters' / f'{outfile_name}.yaml').exists():
        outfile_name = f'{name}_{i:03d}'
        i += 1

    with open(shared.user_data_dir / 'characters' / f'{outfile_name}.yaml', 'w', encoding='utf-8') as f:
        f.write(yaml_data)

    if img is not None:
        img.save(shared.user_data_dir / 'characters' / f'{outfile_name}.png')

    logger.info(f'New character saved to "{shared.user_data_dir}/characters/{outfile_name}.yaml".')
    return gr.update(value=outfile_name, choices=get_available_characters())


def build_pygmalion_style_context(data):
    context = ""
    if 'char_persona' in data and data['char_persona'] != '':
        context += f"{data['char_name']}'s Persona: {data['char_persona']}\n"

    if 'world_scenario' in data and data['world_scenario'] != '':
        context += f"Scenario: {data['world_scenario']}\n"

    if 'example_dialogue' in data and data['example_dialogue'] != '':
        context += f"{data['example_dialogue'].strip()}\n"

    context = f"{context.strip()}\n"
    return context


def upload_tavern_character(img_path, _json):
    _json = {'char_name': _json['name'], 'char_persona': _json['description'], 'char_greeting': _json['first_mes'], 'example_dialogue': _json['mes_example'], 'world_scenario': _json['scenario']}
    return upload_character(json.dumps(_json), img_path, tavern=True)


def check_tavern_character(img_path):
    import gradio as gr
    img = open_image_safely(img_path)

    if img is None:
        return "Invalid or disallowed image file.", None, None, gr.update(interactive=False)

    if "chara" not in img.info:
        return "Not a TavernAI card", None, None, gr.update(interactive=False)

    decoded_string = base64.b64decode(img.info['chara']).replace(b'\\r\\n', b'\\n')
    _json = json.loads(decoded_string)
    if "data" in _json:
        _json = _json["data"]

    return _json['name'], _json['description'], _json, gr.update(interactive=True)


def upload_your_profile_picture(img_path):
    img = open_image_safely(img_path)
    cache_folder = Path(shared.args.disk_cache_dir)
    if not cache_folder.exists():
        cache_folder.mkdir()

    if img is None:
        if Path(f"{cache_folder}/pfp_me.png").exists():
            Path(f"{cache_folder}/pfp_me.png").unlink()
    else:
        img = make_thumbnail(img)
        img.save(Path(f'{cache_folder}/pfp_me.png'))
        logger.info(f'Profile picture saved to "{cache_folder}/pfp_me.png"')


def generate_character_yaml(name, greeting, context):
    data = {
        'name': name,
        'greeting': greeting,
        'context': context,
    }

    data = {k: v for k, v in data.items() if v}  # Strip falsy
    return yaml.dump(data, sort_keys=False, width=float("inf"))


def generate_instruction_template_yaml(instruction_template):
    data = {
        'instruction_template': instruction_template
    }

    return my_yaml_output(data)


def save_character(name, greeting, context, picture, filename):
    filename = sanitize_filename(filename)
    if filename == "":
        logger.error("The filename is empty, so the character will not be saved.")
        return

    data = generate_character_yaml(name, greeting, context)
    filepath = shared.user_data_dir / 'characters' / f'{filename}.yaml'
    save_file(filepath, data)
    path_to_img = shared.user_data_dir / 'characters' / f'{filename}.png'
    if picture is not None:
        # Copy the image file from its source path to the character folder
        shutil.copy(picture, path_to_img)
        logger.info(f'Saved {path_to_img}.')


def delete_character(name, instruct=False):
    name = sanitize_filename(name)
    # Check for character data files
    for extension in ["yml", "yaml", "json"]:
        delete_file(shared.user_data_dir / 'characters' / f'{name}.{extension}')

    # Check for character image files
    for extension in ["png", "jpg", "jpeg"]:
        delete_file(shared.user_data_dir / 'characters' / f'{name}.{extension}')


def generate_user_pfp_cache(user):
    """Generate cached profile picture for user"""
    cache_folder = Path(shared.args.disk_cache_dir)
    if not cache_folder.exists():
        cache_folder.mkdir()

    for path in [shared.user_data_dir / 'users' / f"{user}.{extension}" for extension in ['png', 'jpg', 'jpeg']]:
        if path.exists():
            original_img = Image.open(path)
            # Define file paths
            pfp_path = Path(f'{cache_folder}/pfp_me.png')

            # Save thumbnail
            thumb = make_thumbnail(original_img)
            thumb.save(pfp_path, format='PNG')
            logger.info(f'User profile picture cached to "{pfp_path}"')

            return str(pfp_path)

    return None


def load_user(user_name, name1, user_bio):
    """Load user profile from YAML file"""
    picture = None

    filepath = None
    for extension in ["yml", "yaml", "json"]:
        filepath = shared.user_data_dir / 'users' / f'{user_name}.{extension}'
        if filepath.exists():
            break

    if filepath is None or not filepath.exists():
        logger.error(f"Could not find the user \"{user_name}\" inside {shared.user_data_dir}/users. No user has been loaded.")
        raise ValueError

    with open(filepath, 'r', encoding='utf-8') as f:
        file_contents = f.read()

    extension = filepath.suffix[1:]  # Remove the leading dot
    data = json.loads(file_contents) if extension == "json" else yaml.safe_load(file_contents)

    # Clear existing user picture cache
    cache_folder = Path(shared.args.disk_cache_dir)
    pfp_path = Path(f"{cache_folder}/pfp_me.png")
    if pfp_path.exists():
        pfp_path.unlink()

    # Generate new picture cache
    picture = generate_user_pfp_cache(user_name)

    # Get user name
    if 'name' in data and data['name'] != '':
        name1 = data['name']

    # Get user bio
    if 'user_bio' in data:
        user_bio = data['user_bio']

    return name1, user_bio, picture


def generate_user_yaml(name, user_bio):
    """Generate YAML content for user profile"""
    data = {
        'name': name,
        'user_bio': user_bio,
    }

    return yaml.dump(data, sort_keys=False, width=float("inf"))


def save_user(name, user_bio, picture, filename):
    """Save user profile to YAML file"""
    filename = sanitize_filename(filename)
    if filename == "":
        logger.error("The filename is empty, so the user will not be saved.")
        return

    # Ensure the users directory exists
    users_dir = shared.user_data_dir / 'users'
    users_dir.mkdir(parents=True, exist_ok=True)

    data = generate_user_yaml(name, user_bio)
    filepath = shared.user_data_dir / 'users' / f'{filename}.yaml'
    save_file(filepath, data)

    path_to_img = shared.user_data_dir / 'users' / f'{filename}.png'
    if picture is not None:
        # Copy the image file from its source path to the users folder
        shutil.copy(picture, path_to_img)
        logger.info(f'Saved user profile picture to {path_to_img}.')


def delete_user(name):
    """Delete user profile files"""
    name = sanitize_filename(name)
    # Check for user data files
    for extension in ["yml", "yaml", "json"]:
        delete_file(shared.user_data_dir / 'users' / f'{name}.{extension}')

    # Check for user image files
    for extension in ["png", "jpg", "jpeg"]:
        delete_file(shared.user_data_dir / 'users' / f'{name}.{extension}')


def update_user_menu_after_deletion(idx):
    """Update user menu after a user is deleted"""
    import gradio as gr
    users = get_available_users()
    if len(users) == 0:
        # Create a default user if none exist
        save_user('You', '', None, 'Default')
        users = get_available_users()

    idx = min(int(idx), len(users) - 1)
    idx = max(0, idx)
    return gr.update(choices=users, value=users[idx])


def handle_user_menu_change(state):
    """Handle user menu selection change"""
    try:
        name1, user_bio, picture = load_user(state['user_menu'], state['name1'], state['user_bio'])

        return [
            name1,
            user_bio,
            picture
        ]
    except Exception as e:
        logger.error(f"Failed to load user '{state['user_menu']}': {e}")
        return [
            state['name1'],
            state['user_bio'],
            None
        ]


def handle_save_user_click(name1):
    """Handle save user button click"""
    import gradio as gr
    return [
        name1,
        gr.update(visible=True)
    ]


def my_yaml_output(data):
    '''
    pyyaml is very inconsistent with multiline strings.
    for simple instruction template outputs, this is enough.
    '''
    result = ""
    for k in data:
        result += k + ": |-\n"
        for line in data[k].splitlines():
            result += "  " + line.rstrip(' ') + "\n"

    return result


def handle_send_dummy_message_click(text, state):
    history = send_dummy_message(text, state)
    save_history(history, state['unique_id'], state['character_menu'], state['mode'])
    html = redraw_html(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'])

    return [history, html, {"text": "", "files": []}]


def handle_send_dummy_reply_click(text, state):
    history = send_dummy_reply(text, state)
    save_history(history, state['unique_id'], state['character_menu'], state['mode'])
    html = redraw_html(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'])

    return [history, html, {"text": "", "files": []}]


def handle_remove_last_click(state):
    last_input, history = remove_last_message(state['history'])
    save_history(history, state['unique_id'], state['character_menu'], state['mode'])
    html = redraw_html(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'])

    return [history, html, {"text": last_input, "files": []}]


def handle_unique_id_select(state):
    history = load_history(state['unique_id'], state['character_menu'], state['mode'])
    html = redraw_html(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'])

    # Save this as the last visited chat
    save_last_chat_state(state['character_menu'], state['mode'], state['unique_id'])

    convert_to_markdown.cache_clear()

    return [history, html]


def handle_start_new_chat_click(state):
    import gradio as gr
    history = start_new_chat(state)
    histories = find_all_histories_with_first_prompts(state)
    html = redraw_html(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'])

    convert_to_markdown.cache_clear()

    if len(histories) > 0:
        past_chats_update = gr.update(choices=histories, value=histories[0][1])
    else:
        past_chats_update = gr.update(choices=histories)

    return [history, html, past_chats_update]


def handle_start_incognito_chat_click(state):
    import gradio as gr
    unique_id = 'incognito-' + datetime.now().strftime('%Y%m%d-%H-%M-%S')
    history = start_new_chat(state, unique_id=unique_id)
    html = redraw_html(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'])

    convert_to_markdown.cache_clear()

    histories = find_all_histories_with_first_prompts(state)
    past_chats_update = gr.update(choices=histories, value=unique_id)

    return [history, html, past_chats_update]


def handle_delete_chat_confirm_click(state):
    filtered_histories = find_all_histories_with_first_prompts(state)
    filtered_ids = [h[1] for h in filtered_histories]

    if state['unique_id'] not in filtered_ids:
        # Incognito or unknown chat — just load the most recent saved chat
        index = '0'
    else:
        index = str(filtered_ids.index(state['unique_id']))

    delete_history(state['unique_id'], state['character_menu'], state['mode'])
    history, unique_id = load_history_after_deletion(state, index)
    html = redraw_html(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'])

    convert_to_markdown.cache_clear()

    return [history, html, unique_id]


def handle_branch_chat_click(state):
    import gradio as gr
    branch_from_index = state['branch_index']
    if branch_from_index == -1:
        history = state['history']
    else:
        history = state['history']
        history['visible'] = history['visible'][:branch_from_index + 1]
        history['internal'] = history['internal'][:branch_from_index + 1]
        # Prune the metadata dictionary to remove entries beyond the branch point
        if 'metadata' in history:
            history['metadata'] = {k: v for k, v in history['metadata'].items() if int(k.split('_')[-1]) <= branch_from_index}

    prefix = 'incognito-' if state['unique_id'] and state['unique_id'].startswith('incognito-') else ''
    new_unique_id = prefix + datetime.now().strftime('%Y%m%d-%H-%M-%S')
    save_history(history, new_unique_id, state['character_menu'], state['mode'])

    histories = find_all_histories_with_first_prompts(state)
    html = redraw_html(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'])

    convert_to_markdown.cache_clear()

    past_chats_update = gr.update(choices=histories, value=new_unique_id)

    return [history, html, past_chats_update, -1]


def handle_edit_message_click(state):
    history = state['history']
    message_index = int(state['edit_message_index'])
    new_text = state['edit_message_text']
    role = state['edit_message_role']  # "user" or "assistant"

    if message_index >= len(history['internal']):
        html_output = redraw_html(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'])
        return [history, html_output]

    role_idx = 0 if role == "user" else 1

    if 'metadata' not in history:
        history['metadata'] = {}

    key = f"{role}_{message_index}"
    if key not in history['metadata']:
        history['metadata'][key] = {}

    # If no versions exist yet for this message, store the current (pre-edit) content as the first version.
    if "versions" not in history['metadata'][key] or not history['metadata'][key]["versions"]:
        original_content = history['internal'][message_index][role_idx]
        original_visible = history['visible'][message_index][role_idx]
        original_timestamp = history['metadata'][key].get('timestamp', get_current_timestamp())

        version_entry = {
            "content": original_content,
            "visible_content": original_visible,
            "timestamp": original_timestamp
        }
        ts = history['metadata'][key].get('tool_sequence')
        if ts is not None:
            version_entry['tool_sequence'] = ts
        history['metadata'][key]["versions"] = [version_entry]

    history['internal'][message_index][role_idx] = apply_extensions('input', new_text, state, is_chat=True)
    history['visible'][message_index][role_idx] = html.escape(new_text)
    history['metadata'][key].pop('tool_sequence', None)

    add_message_version(history, role, message_index, is_current=True)

    save_history(history, state['unique_id'], state['character_menu'], state['mode'])
    html_output = redraw_html(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'])

    return [history, html_output]


def handle_navigate_version_click(state):
    history = state['history']
    message_index = int(state['navigate_message_index'])
    direction = state['navigate_direction']
    role = state['navigate_message_role']

    if not role:
        logger.error("Role not provided for version navigation.")
        html = redraw_html(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'])
        return [history, html]

    key = f"{role}_{message_index}"
    if 'metadata' not in history or key not in history['metadata'] or 'versions' not in history['metadata'][key]:
        html = redraw_html(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'])
        return [history, html]

    metadata = history['metadata'][key]
    versions = metadata['versions']
    # Default to the last version if current_version_index is not set
    current_idx = metadata.get('current_version_index', len(versions) - 1 if versions else 0)

    if direction == 'left':
        new_idx = max(0, current_idx - 1)
    else:  # right
        new_idx = min(len(versions) - 1, current_idx + 1)

    if new_idx == current_idx:
        html = redraw_html(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'])
        return [history, html]

    msg_content_idx = 0 if role == 'user' else 1  # 0 for user content, 1 for assistant content in the pair
    version_to_load = versions[new_idx]
    history['internal'][message_index][msg_content_idx] = version_to_load['content']
    history['visible'][message_index][msg_content_idx] = version_to_load['visible_content']
    metadata['current_version_index'] = new_idx

    # Restore per-version tool_sequence so follow-up prompts see consistent context
    version_ts = version_to_load.get('tool_sequence')
    if version_ts is not None:
        metadata['tool_sequence'] = version_ts
    else:
        metadata.pop('tool_sequence', None)

    update_message_metadata(history['metadata'], role, message_index, timestamp=version_to_load['timestamp'])

    # Redraw and save
    html = redraw_html(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'])
    save_history(history, state['unique_id'], state['character_menu'], state['mode'])

    return [history, html]


def handle_rename_chat_click():
    import gradio as gr
    return [
        gr.update(value="My New Chat"),
        gr.update(visible=True),
    ]


def handle_rename_chat_confirm(rename_to, state):
    import gradio as gr

    if state['unique_id'] and state['unique_id'].startswith('incognito-'):
        return [
            gr.update(),
            gr.update(visible=False),
        ]

    rename_history(state['unique_id'], rename_to, state['character_menu'], state['mode'])
    histories = find_all_histories_with_first_prompts(state)

    return [
        gr.update(choices=histories, value=rename_to),
        gr.update(visible=False),
    ]


def handle_search_chat_change(state):
    import gradio as gr
    histories = find_all_histories_with_first_prompts(state)
    return gr.update(choices=histories)


def handle_upload_chat_history(load_chat_history, state):
    import gradio as gr
    history = start_new_chat(state)
    history = load_history_json(load_chat_history, history)
    save_history(history, state['unique_id'], state['character_menu'], state['mode'])
    histories = find_all_histories_with_first_prompts(state)

    html = redraw_html(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'])

    convert_to_markdown.cache_clear()

    if len(histories) > 0:
        past_chats_update = gr.update(choices=histories, value=histories[0][1])
    else:
        past_chats_update = gr.update(choices=histories)

    return [
        history,
        html,
        past_chats_update
    ]


def handle_character_menu_change(state):
    import gradio as gr
    name1, name2, picture, greeting, context = load_character(state['character_menu'], state['name1'], state['name2'])

    state['name1'] = name1
    state['name2'] = name2
    state['character_picture'] = picture
    state['greeting'] = greeting
    state['context'] = context

    history, loaded_unique_id = load_latest_history(state)
    histories = find_all_histories_with_first_prompts(state)
    html = redraw_html(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'])

    convert_to_markdown.cache_clear()

    if len(histories) > 0:
        past_chats_update = gr.update(choices=histories, value=loaded_unique_id or histories[0][1])
    else:
        past_chats_update = gr.update(choices=histories)

    return [
        history,
        html,
        name1,
        name2,
        picture,
        greeting,
        context,
        past_chats_update
    ]


def handle_character_picture_change(picture_path):
    """Update or clear cache when character picture changes"""
    picture = open_image_safely(picture_path)
    cache_folder = Path(shared.args.disk_cache_dir)
    if not cache_folder.exists():
        cache_folder.mkdir()

    if picture is not None:
        # Save to cache
        picture.save(Path(f'{cache_folder}/pfp_character.png'), format='PNG')
        thumb = make_thumbnail(picture)
        thumb.save(Path(f'{cache_folder}/pfp_character_thumb.png'), format='PNG')
    else:
        # Remove cache files when picture is cleared
        for cache_file in ['pfp_character.png', 'pfp_character_thumb.png']:
            cache_path = Path(f'{cache_folder}/{cache_file}')
            if cache_path.exists():
                cache_path.unlink()


def handle_mode_change(state):
    import gradio as gr
    history, loaded_unique_id = load_latest_history(state)
    histories = find_all_histories_with_first_prompts(state)

    # Ensure character picture cache exists
    if state['mode'] in ['chat', 'chat-instruct'] and state['character_menu'] and state['character_menu'] != 'None':
        generate_pfp_cache(state['character_menu'])

    html = redraw_html(history, state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'])

    convert_to_markdown.cache_clear()

    if len(histories) > 0:
        past_chats_update = gr.update(choices=histories, value=loaded_unique_id or histories[0][1])
    else:
        past_chats_update = gr.update(choices=histories)

    return [
        history,
        html,
        gr.update(visible=state['mode'] != 'instruct'),
        gr.update(visible=state['mode'] == 'chat-instruct'),
        past_chats_update
    ]


def handle_save_character_click(name2):
    import gradio as gr
    return [
        name2,
        gr.update(visible=True)
    ]


def handle_load_template_click(instruction_template):
    from modules.models_settings import load_instruction_template
    output = load_instruction_template(instruction_template)
    return [
        output,
        "Select template to load..."
    ]


def handle_save_template_click(instruction_template_str):
    import gradio as gr
    contents = generate_instruction_template_yaml(instruction_template_str)
    root = str(shared.user_data_dir / 'instruction-templates') + '/'
    return [
        "My Template.yaml",
        root,
        contents,
        root,
        gr.update(visible=True)
    ]


def handle_delete_template_click(template):
    import gradio as gr
    root = str(shared.user_data_dir / 'instruction-templates') + '/'
    return [
        f"{template}.yaml",
        root,
        root,
        gr.update(visible=False)
    ]


def handle_your_picture_change(picture, state):
    upload_your_profile_picture(picture)
    html = redraw_html(state['history'], state['name1'], state['name2'], state['mode'], state['chat_style'], state['character_menu'], reset_cache=True)

    return html


def handle_send_instruction_click(state):
    import gradio as gr
    state['mode'] = 'instruct'
    state['history'] = {'internal': [], 'visible': [], 'metadata': {}}

    output = generate_chat_prompt("Input", state)

    if state["show_two_notebook_columns"]:
        return gr.update(), output, ""
    else:
        return output, gr.update(), gr.update()


def handle_send_chat_click(state):
    import gradio as gr
    output = generate_chat_prompt("", state, _continue=True)

    if state["show_two_notebook_columns"]:
        return gr.update(), output, ""
    else:
        return output, gr.update(), gr.update()
